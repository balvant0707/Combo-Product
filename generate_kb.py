"""
MixBox Knowledge Base — Word Document Generator
Reads 8 markdown pages from docs/knowledge-base/ and generates a
fully-styled .docx with embedded screenshots.

Screenshots expected in: docs/knowledge-base/screenshots/
  01-dashboard.png
  02-growth-apps.png
  03-choose-bundle-type.png
  04-create-simple-box.png
  05-create-specific-box.png
  06-steps-configuration.png
  07-manage-boxes.png
  08-analytics.png
  09-analytics-charts.png
  10-widget-settings.png
"""

import os, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ──────────────────────────────────────────────────────────────────────
BASE        = r"c:\shopify apps\combo-product"
KB_DIR      = os.path.join(BASE, "docs", "knowledge-base")
SS_DIR      = os.path.join(KB_DIR, "screenshots")
OUT_PATH    = os.path.join(BASE, "MixBox_Knowledge_Base.docx")

# ── Brand colours ──────────────────────────────────────────────────────────────
GREEN   = RGBColor(0x2A, 0x7A, 0x4F)
DARK    = RGBColor(0x1E, 0x1E, 0x2D)
GRAY    = RGBColor(0x6B, 0x72, 0x80)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
BLUE    = RGBColor(0x1D, 0x4E, 0xD8)
RED     = RGBColor(0xDC, 0x26, 0x26)
AMBER   = RGBColor(0xD9, 0x77, 0x06)
LGRAY_H = "F3F4F6"
GREEN_H = "2A7A4F"

# ── Helpers ────────────────────────────────────────────────────────────────────

def _font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name   = name
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _shd(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def para(doc, text="", bold=False, italic=False, color=None, size=11,
         indent=0, bullet=False, center=False, space_before=2, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if bullet:
        p.style = doc.styles["List Bullet"]
    if indent:
        p.paragraph_format.left_indent = Inches(indent * 0.3)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text:
        r = p.add_run(text)
        _font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20 if level == 1 else 14 if level == 2 else 10)
    p.paragraph_format.space_after  = Pt(6)
    sizes  = {1: 20, 2: 15, 3: 13}
    colors = {1: GREEN, 2: DARK, 3: DARK}
    r = p.add_run(text)
    _font(r, size=sizes.get(level, 12), bold=True, color=colors.get(level, DARK))
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "6")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "2A7A4F")
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def note_box(doc, text, kind="tip"):
    """Coloured callout paragraph: tip / warning / important."""
    colors_map = {"tip": ("E6F4EA", GREEN), "warning": ("FFF8E1", AMBER), "important": ("FFEBEE", RED)}
    bg, fg = colors_map.get(kind, ("F3F4F6", DARK))
    labels  = {"tip": "TIP", "warning": "WARNING", "important": "IMPORTANT"}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Inches(0.1)
    r1 = p.add_run(f"  {labels[kind]}: ")
    _font(r1, bold=True, size=10, color=fg)
    r2 = p.add_run(text)
    _font(r2, size=10, color=DARK, italic=True)
    return p


def table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style     = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    # header row
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].font.bold  = True
        c.paragraphs[0].runs[0].font.color.rgb = WHITE
        c.paragraphs[0].runs[0].font.size  = Pt(10)
        _shd(c, GREEN_H)
    # data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = tbl.rows[ri + 1].cells[ci]
            c.text = str(val)
            c.paragraphs[0].runs[0].font.size = Pt(9.5)
            if ri % 2 == 0:
                _shd(c, LGRAY_H)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                if i < len(row.cells):
                    row.cells[i].width = Inches(w)
    doc.add_paragraph()   # spacing after table
    return tbl


def screenshot(doc, filename, caption=""):
    """Insert a screenshot image centred with an optional caption."""
    path = os.path.join(SS_DIR, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run()
        run.add_picture(path, width=Inches(6.0))
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cap.add_run(f"↑  {caption}")
            _font(r, size=9, italic=True, color=GRAY)
            cap.paragraph_format.space_after = Pt(10)
    else:
        # placeholder when screenshot not yet placed
        p = doc.add_paragraph()
        r = p.add_run(f"[ Screenshot: {filename} ]")
        _font(r, size=9.5, italic=True, color=GRAY)
        p.paragraph_format.space_after = Pt(4)
        if caption:
            cap = doc.add_paragraph()
            r2 = cap.add_run(caption)
            _font(r2, size=9, italic=True, color=GRAY)


def checklist(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.left_indent  = Inches(0.2)
        r = p.add_run(f"☐  {item}")
        _font(r, size=10.5)


def field_row(doc, label, ftype, desc, required=False, default=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.left_indent  = Inches(0.25)
    r1 = p.add_run(label)
    _font(r1, bold=True, size=10.5, color=DARK)
    r2 = p.add_run(f"  [{ftype}]")
    _font(r2, size=9.5, italic=True, color=BLUE)
    if required:
        r3 = p.add_run("  *Required")
        _font(r3, size=9, color=RED)
    if default:
        r4 = p.add_run(f"  Default: {default}")
        _font(r4, size=9, italic=True, color=GRAY)
    r5 = p.add_run(f"\n    → {desc}")
    _font(r5, size=10, color=GRAY)


# ══════════════════════════════════════════════════════════════════════════════
#  BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()

# page setup
sec = doc.sections[0]
sec.page_width    = Inches(8.5)
sec.page_height   = Inches(11)
sec.left_margin   = Inches(1)
sec.right_margin  = Inches(1)
sec.top_margin    = Inches(1)
sec.bottom_margin = Inches(1)

# ── Cover Page ─────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(90)
r = p.add_run("MixBox")
_font(r, size=38, bold=True, color=GREEN)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Box & Bundle Builder")
_font(r2, size=20, color=DARK)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Knowledge Base")
_font(r3, size=26, bold=True, color=DARK)

doc.add_paragraph()
p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run("Complete Guide for Shopify Merchants")
_font(r4, size=13, italic=True, color=GRAY)

doc.add_paragraph()
p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = p5.add_run("Version 1.0  |  April 2026  |  Pryxotech")
_font(r5, size=10, color=GRAY)

doc.add_page_break()

# ── Table of Contents ──────────────────────────────────────────────────────────
heading(doc, "Table of Contents", 1)
toc = [
    ("Page 1", "Getting Started"),
    ("Page 2", "Dashboard Overview"),
    ("Page 3", "Create a Simple Box"),
    ("Page 4", "Create a Specific Box"),
    ("Page 5", "Manage Boxes"),
    ("Page 6", "Analytics"),
    ("Page 7", "Widget Settings"),
    ("Page 8", "FAQ & Troubleshooting"),
]
for num, title in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{num}  ")
    _font(r1, bold=True, size=11, color=GREEN)
    r2 = p.add_run(title)
    _font(r2, size=11, color=DARK)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  PAGE 1 — GETTING STARTED
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "Page 1 — Getting Started", 1)
para(doc, "Get your first combo bundle live on your storefront in under 10 minutes.", italic=True, color=GRAY)

heading(doc, "What is MixBox?", 2)
para(doc, "MixBox lets you create combo boxes and bundle builders directly inside your Shopify store. "
         "Customers can pick products to fill a box and add the whole bundle to their cart in one click — "
         "great for gift sets, starter kits, meal preps, beauty bundles, and more.")

heading(doc, "Two Bundle Types", 3)
table(doc,
    ["Type", "Best For"],
    [
        ["Simple Box",   "One product grid — customer picks any N items from your store or a category"],
        ["Specific Box", "Multi-step wizard where each step shows a different product group (e.g. Step 1 = base, Step 2 = topping)"],
    ],
    [1.8, 4.4]
)

heading(doc, "Step 1 — Install & Open the App", 2)
for s in [
    "Install MixBox from the Shopify App Store.",
    "Open your Shopify Admin → Apps → MixBox.",
    "You will land on the Dashboard (your home screen).",
]:
    para(doc, s, bullet=True, size=10.5)

heading(doc, "Step 2 — Enable the App Embed", 2)
para(doc, "Before customers can see your bundles, you need to turn on the app embed in your theme.")
for s in [
    "On the Dashboard, find the 'Theme App Embed Status' card (bottom-left).",
    "If it shows Off, click the card or use the 'Open Shopify Theme Editor' button.",
    "In the Theme Editor, find MixBox – Box & Bundle Builder under the Apps section and toggle it On.",
    "Click Save in the theme editor.",
]:
    para(doc, s, bullet=True, size=10.5)
note_box(doc, "If the embed is Off, no bundles will appear on your storefront — even if you have created boxes.", "important")

heading(doc, "Step 3 — Add the Widget to Your Theme", 2)
for s in [
    "Click 'Open Shopify Theme Editor' on the Dashboard.",
    "Navigate to the template where you want the widget (e.g. a Product page or a custom page).",
    "In the left sidebar, click 'Add section' → Apps → Combo Builder.",
    "Drag the block to the position you want on the page.",
    "Click Save.",
]:
    para(doc, s, bullet=True, size=10.5)
note_box(doc, "A dedicated 'Build Your Bundle' landing page often converts very well.", "tip")

heading(doc, "Step 4 — Create Your First Box", 2)
for s in [
    "From the Dashboard, click 'Create Bundle Box' (or go to Manage Boxes → + Create Box).",
    "Choose Simple Box or Specific Box.",
    "Fill in the required fields (Title, Items Required, Price).",
    "Click Save & Publish.",
]:
    para(doc, s, bullet=True, size=10.5)

heading(doc, "Quick-Start Checklist", 2)
checklist(doc, [
    "App installed and opened",
    "Theme App Embed turned On",
    "Combo Builder block added to your theme",
    "First box created and published",
    "Visited your storefront to confirm the widget appears",
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  PAGE 2 — DASHBOARD
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "Page 2 — Dashboard Overview", 1)
para(doc, "Understand every number and card on your home screen at a glance.", italic=True, color=GRAY)

screenshot(doc, "01-dashboard.png", "Dashboard home screen — stats, setup cards, quick actions, recent orders")

heading(doc, "Top Stats Bar", 2)
para(doc, "Five metric cards across the top summarise your store's bundle activity.")
table(doc,
    ["Card", "What It Shows", "Notes"],
    [
        ["Live",              "Boxes currently published and visible to customers",           "Only active boxes count"],
        ["Order Credit Left", "Remaining bundle orders allowed this month",                   "'0/2 used' = 0 used of 2 total"],
        ["Orders",            "Total bundle orders placed in the last 30 days",              "Links to Analytics page"],
        ["Total Revenue",     "Revenue from bundles in the last 30 days",                    "In your store currency"],
        ["Conversion Rate",   "% of widget visitors who completed a bundle purchase",        "Last 30 days"],
    ],
    [1.6, 2.8, 1.8]
)
note_box(doc, "If Order Credit Left is running low, upgrade your plan before you hit the limit — orders above the limit will not be processed.", "warning")

heading(doc, "Theme App Embed Status", 2)
para(doc, "Shows whether the MixBox widget is enabled in your active theme.")
table(doc,
    ["Status", "Meaning", "Action"],
    [
        ["🟢 On",  "Widget can appear on storefront",    "No action needed"],
        ["⚪ Off", "Widget is disabled for all customers", "Click the card to open Theme Editor and toggle On"],
    ],
    [1.0, 3.0, 2.2]
)

heading(doc, "Theme Widget Setup", 2)
para(doc, "4-step visual guide to placing the widget on your storefront:")
for i, s in enumerate([
    "Opens Theme Customization on your live product template.",
    "Combo Builder block is auto-added to the Apps section.",
    "Drag the block to the right position.",
    "Click Save — your storefront is live.",
], 1):
    para(doc, f"{i}.  {s}", indent=1, size=10.5)

heading(doc, "Quick Actions", 2)
table(doc,
    ["Button", "Where It Goes"],
    [
        ["Create Bundle Box", "Opens the 'Choose Bundle Type' modal to start a new box"],
        ["Manage Boxes",      "Takes you to the full list of all your boxes"],
        ["View Analytics",    "Opens the Analytics page"],
        ["Widget Settings",   "Opens the appearance and behaviour settings"],
    ],
    [2.0, 4.2]
)

heading(doc, "Recent Orders", 2)
para(doc, "Shows the latest bundle orders from your store. If no orders exist yet, you will see: "
         "'No combo box orders yet.' Click View all (top-right) to open the full orders list in Analytics.")

screenshot(doc, "02-growth-apps.png", "Recommended Growth Apps section — CartLift and Fomoify")

heading(doc, "Recommended Growth Apps", 2)
table(doc,
    ["App", "Purpose"],
    [
        ["CartLift: Cart Drawer and Upsell", "Grow average order value with cart drawer upsells and smart cart offers"],
        ["Fomoify Sales Popup and Proof",    "Build trust with real-time sales popups and social proof nudges"],
    ],
    [2.5, 3.7]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  PAGE 3 — CREATE SIMPLE BOX
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "Page 3 — Create a Simple Box", 1)
para(doc, "Set up a 'fill the box' bundle where customers pick any N products from your store.", italic=True, color=GRAY)

screenshot(doc, "03-choose-bundle-type.png", "Choose Bundle Type modal — Simple Box vs Specific Box")

heading(doc, "What is a Simple Box?", 2)
para(doc, "A Simple Box is a preconfigured bundle where you decide how many items the customer must pick, "
         "they browse your catalogue, and the whole selection is added to the cart at a fixed or dynamic price.")
para(doc, "Best used for: snack boxes, gift hampers, beauty kits, subscription starter packs — "
         "any bundle where all products come from the same general catalogue.", italic=True, color=GRAY, size=10)

heading(doc, "How to Open the Simple Box Creator", 2)
for s in [
    "From the Dashboard, click 'Create Bundle Box'.",
    "In the 'Choose Bundle Type' pop-up, click 'Create Box' under Create Simple Box.",
]:
    para(doc, s, bullet=True, size=10.5)

screenshot(doc, "04-create-simple-box.png", "Create Simple Box — General Configuration form")

heading(doc, "Status Toggle", 2)
table(doc,
    ["Status", "Meaning"],
    [
        ["On (default)", "Box is visible to customers as soon as you save"],
        ["Off",          "Box is saved as a draft — customers cannot see it until you turn it On"],
    ],
    [1.5, 4.7]
)
note_box(doc, "Build and preview your box with Status Off, then switch it On when you are happy.", "tip")

heading(doc, "General Configuration Fields", 2)
field_row(doc, "Title",           "Text",   "Customer-facing name of the bundle. E.g. 'Build Your Perfect Snack Box'", required=True)
field_row(doc, "CTA Button Text", "Text",   "Text on the main call-to-action button on the storefront.", default="Build Your Own Box")
field_row(doc, "Add Button Text", "Text",   "Text on the final button that adds the bundle to the cart.", default="Add To Cart")
field_row(doc, "Items Required",  "Number", "How many products the customer must select (2–8).", required=True)

heading(doc, "Pricing", 2)
table(doc,
    ["Option", "How It Works", "Best For"],
    [
        ["Fixed Price",   "You set one flat price for the bundle regardless of which products are picked",     "Gift boxes with a clear price point"],
        ["Dynamic Price", "Bundle price = sum of the individual product prices selected by the customer",      "Mix-and-match catalogues with variable pricing"],
    ],
    [1.4, 3.0, 1.8]
)

heading(doc, "Image Upload", 2)
para(doc, "Upload a cover image for your bundle. Accepted: JPG, PNG, WEBP, GIF, AVIF. Maximum size: 5 MB. "
         "This image appears in the widget header on the storefront.")

heading(doc, "Choose Display Scope", 2)
table(doc,
    ["Option", "What Customers Can Pick From"],
    [
        ["Whole Store",          "Every product in your store (default)"],
        ["Specific Collections", "Only products from collections you choose"],
    ],
    [2.0, 4.2]
)

heading(doc, "Options", 2)
field_row(doc, "Enable Gift Box Option",     "Toggle", "Show a gift wrapping option to customers at checkout.", default="Off")
field_row(doc, "Enable Gift Message Field",  "Toggle", "Display a text area for a personal gift message, saved with the order.", default="Off")
field_row(doc, "Allow Duplicate Products",   "Toggle", "Let customers pick the same product more than once to fill multiple slots.", default="Off")

heading(doc, "Example Use Case", 2)
para(doc, "A bakery creates 'Build Your Biscuit Box' with 6 items required, Fixed Price ₹799, "
         "Whole Store scope, and Allow Duplicates On. Customers pick 6 biscuit packets and add the "
         "₹799 bundle to cart.", italic=True, size=10.5)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  PAGE 4 — CREATE SPECIFIC BOX
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "Page 4 — Create a Specific Box", 1)
para(doc, "Build a guided, multi-step bundle where each step offers a different curated set of products.", italic=True, color=GRAY)

heading(doc, "What is a Specific Box?", 2)
para(doc, "A Specific Box is a step-by-step bundle builder. You divide the selection into multiple Steps — "
         "each step shows its own group of products. Customers walk through each step on the storefront "
         "and pick one product per step.")
para(doc, "Best used for: makeup kits, meal prep boxes, skincare routines, tech bundles — "
         "any bundle where each slot must come from a specific category.", italic=True, color=GRAY, size=10)

screenshot(doc, "05-create-specific-box.png", "Create Specific Box — General Configuration and Price Type")

heading(doc, "General Configuration Fields", 2)
field_row(doc, "Title",               "Text",   "Storefront name of your bundle.", required=True)
field_row(doc, "Description",         "Text",   "Short sentence explaining the bundle, e.g. 'Build Your Own Makeup Kit'.")
field_row(doc, "Bundle Button Text",  "Text",   "Text on the main CTA button.", default="BUILD YOUR OWN BOX")
field_row(doc, "Steps",               "Number", "Number of selection steps (2–8). Use – and + buttons to adjust.", required=True)
field_row(doc, "Image",               "Upload", "Cover image for the bundle. Max 2 MB. Types: JPG/PNG/WEBP/GIF/AVIF.")

heading(doc, "Price Type", 2)
table(doc,
    ["Option", "How Price Is Calculated", "Discount Option Available?"],
    [
        ["Fixed Price",   "One flat price you set for the entire bundle",          "No"],
        ["Dynamic Price", "Sum of all selected products across all steps",         "Yes — Percentage or Fixed Amount"],
    ],
    [1.5, 3.0, 1.7]
)

heading(doc, "Discount Type (Dynamic Price only)", 2)
table(doc,
    ["Option", "Effect"],
    [
        ["None",           "No discount — customer pays the full product sum"],
        ["Percentage (%)", "Apply a flat percentage discount to the total (e.g. 10% off)"],
        ["Fixed Amount",   "Subtract a fixed amount from the total (e.g. ₹100 off)"],
    ],
    [1.5, 4.7]
)

heading(doc, "Additional Options", 2)
field_row(doc, "Enable Gift Packaging Option", "Toggle", "Offer gift wrapping to customers.", default="Off")
field_row(doc, "Enable Gift Note Field",        "Toggle", "Show a text box for a personal message.", default="Off")
field_row(doc, "Allow Repeating Products",      "Toggle", "Let customers select the same product in more than one step.", default="Off")

screenshot(doc, "06-steps-configuration.png", "Steps Configuration — Step Product Picker and Step Content Settings")

heading(doc, "Steps Configuration", 2)
para(doc, "Each step you defined gets its own configuration card. Click the step tabs (Step 1, Step 2 …) to switch between them.")

heading(doc, "Step Product Picker Setup", 3)
field_row(doc, "Step Name",                "Text",        "Heading customers see for this step. E.g. 'Choose Your Protein'.", required=True)
field_row(doc, "Step Scope",               "Select",      "Which products appear in this step: Select Collections or Whole Store.")
field_row(doc, "Choose Step Collections",  "Button",      "Opens a collection picker. Selected count shown as a badge (e.g. '3 selected').")

heading(doc, "Step Content Settings", 3)
field_row(doc, "Step Heading",             "Text",        "Large heading above the product grid for this step.", default="Choose Your Main Product")
field_row(doc, "Step Description",         "Text",        "Subtext guiding the customer for this step.")
field_row(doc, "Step Selection Button Text","Text",       "Button text on each product card.", default="Confirm Selection")
field_row(doc, "Make This Step Optional",  "Toggle",      "Customers can skip this step. Good for bonus add-ons.", default="Off")

note_box(doc, "Mark the last step as Optional if you want to include a 'bonus add-on' that customers can choose or skip.", "tip")

heading(doc, "Example Use Case", 2)
para(doc, "A cosmetics brand creates 'Build Your Beauty Kit' with 3 steps: "
         "Step 1 = Choose Foundation (scoped to Foundation collection), "
         "Step 2 = Choose Lipstick (Lipstick collection), "
         "Step 3 = Choose Mascara — optional (Eye collection). "
         "Price = Dynamic with 10% discount.", italic=True, size=10.5)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  PAGE 5 — MANAGE BOXES
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "Page 5 — Manage Boxes", 1)
para(doc, "Find, filter, edit, activate, or delete your bundle boxes from one central place.", italic=True, color=GRAY)

screenshot(doc, "07-manage-boxes.png", "Manage Boxes page — stats cards, search bar, filter tabs, empty state")

heading(doc, "Stats Cards", 2)
table(doc,
    ["Card", "What It Shows"],
    [
        ["Total Boxes",    "All boxes ever created (active + inactive)"],
        ["Active Boxes",   "Boxes currently published and visible to customers"],
        ["Inactive Boxes", "Boxes saved as drafts or hidden from customers"],
        ["Total Orders",   "All-time bundle orders across every box"],
    ],
    [2.0, 4.2]
)

heading(doc, "Searching and Filtering", 2)
field_row(doc, "Search Bar",      "Text Input",   "Type any part of a box name to filter results instantly.")
field_row(doc, "All Tab",         "Filter Tab",   "Shows every box (active + inactive) with total count.")
field_row(doc, "Active Tab",      "Filter Tab",   "Shows only currently live boxes.")
field_row(doc, "Inactive Tab",    "Filter Tab",   "Shows only hidden/draft boxes.")

note_box(doc, "Use the Inactive filter to find boxes you paused and want to re-activate for a sale or season.", "tip")

heading(doc, "Actions You Can Take", 2)
table(doc,
    ["Action", "How To", "Effect"],
    [
        ["Edit",       "Click Edit button or the box name",   "Opens full edit form — all settings can be changed"],
        ["Activate",   "Toggle Status switch On",             "Box goes live on storefront immediately"],
        ["Deactivate", "Toggle Status switch Off",            "Box is hidden from storefront but data is preserved"],
        ["Duplicate",  "Click Duplicate",                     "Creates a copy of the box as a new draft"],
        ["Delete",     "Click Delete → confirm",              "Permanently removes box AND its Shopify product"],
    ],
    [1.2, 2.0, 3.0]
)

note_box(doc, "Delete is permanent and also removes the linked Shopify product. Use Deactivate if you just want to hide the box temporarily.", "warning")

heading(doc, "Create Box Button", 2)
para(doc, "The '+ Create Box' button (top-right) opens the Choose Bundle Type modal, "
         "letting you create a new Simple or Specific box at any time.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  PAGE 6 — ANALYTICS
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "Page 6 — Analytics", 1)
para(doc, "Track how your bundles are performing, which products are most popular, and how revenue trends over time.", italic=True, color=GRAY)

screenshot(doc, "08-analytics.png", "Analytics page — Performance Overview with filters and metric cards")

heading(doc, "Performance Overview", 2)

heading(doc, "Type Filter", 3)
table(doc,
    ["Filter", "Shows Data For"],
    [
        ["All Box",       "Combined data across both Simple and Specific boxes"],
        ["Simple Box",    "Simple boxes only"],
        ["Specific Box",  "Specific boxes only"],
    ],
    [1.8, 4.4]
)

heading(doc, "Date Range Picker", 3)
para(doc, "Pre-set options: Last 7 Days, Last 30 Days (default), Last 90 Days, Custom Range.")
para(doc, "The date bar shows the exact current and previous period — e.g. Current: 21 Mar – 20 Apr vs "
         "Previous: 19 Feb – 21 Mar. This lets you instantly compare performance period-over-period.")

heading(doc, "Sync Orders Button", 3)
para(doc, "Click '↻ Sync Orders' to manually pull the latest order data from Shopify. "
         "Use this if recent orders are not showing up yet.")

heading(doc, "Key Metric Cards", 2)
table(doc,
    ["Card", "What It Shows"],
    [
        ["Total All Revenue",        "Sum of revenue from all bundle orders in the selected period"],
        ["Total All Sold",           "Number of bundle orders completed in the period"],
        ["Average All Order Value",  "Revenue ÷ Orders — your average bundle sale value"],
        ["Active All",               "Number of boxes currently live on your storefront"],
    ],
    [2.2, 4.0]
)

heading(doc, "Most Picked Products", 2)
para(doc, "A ranked list showing which individual products customers selected most often across all bundle orders. "
         "Use this to identify best sellers and ensure popular products stay in stock.")

heading(doc, "Box Performance", 2)
para(doc, "A ranked table comparing each box side-by-side: box name, number of orders, and revenue generated.")

screenshot(doc, "09-analytics-charts.png", "Analytics — Revenue Over Time, Orders Over Time, and Recent Orders table")

heading(doc, "Revenue Over Time Chart", 2)
para(doc, "A line chart showing daily revenue across your selected date range. "
         "Solid line = current period, Dashed line = previous period for comparison.")

heading(doc, "Orders Over Time Chart", 2)
para(doc, "A line chart showing the number of orders placed each day. Same two-line current/previous format.")

heading(doc, "Recent All Orders", 2)
para(doc, "A paginated table of the most recent bundle orders in the selected period. "
         "Use ‹ and › arrows at the bottom-right to move between pages.")

heading(doc, "Tips for Using Analytics", 2)
for tip in [
    "Check weekly — reviewing analytics once a week helps you spot trends early.",
    "Compare periods — use the previous-period dashed line to see improvement or decline.",
    "Most Picked Products — use this list when restocking. If a product sells out it breaks the bundle experience.",
    "Low-performing boxes — if a box has zero orders after 2 weeks, consider changing its title, price, or collections.",
]:
    para(doc, tip, bullet=True, size=10.5)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  PAGE 7 — WIDGET SETTINGS
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "Page 7 — Widget Settings", 1)
para(doc, "Customise how the combo builder widget looks and behaves on your storefront.", italic=True, color=GRAY)

screenshot(doc, "10-widget-settings.png", "Widget Settings — Theme Customizer, Widget Width, Product Grid, Display Options")

heading(doc, "Theme Customizer", 2)
para(doc, "Personalise the colours of the widget to match your brand. Changes apply to all boxes globally.")

heading(doc, "Colour Presets", 3)
table(doc,
    ["Preset", "Colours"],
    [
        ["Forest", "Dark green + green"],
        ["Ocean",  "Dark blue + teal"],
        ["Sunset", "Orange + brown"],
        ["Plum",   "Purple + light purple"],
        ["Rose",   "Pink + red"],
    ],
    [1.5, 4.7]
)

heading(doc, "Custom Colours", 3)
field_row(doc, "Primary Color",   "Hex Input / Colour Picker", "Main brand colour used for buttons, highlights, and active states.", default="#2A7A4F")
field_row(doc, "Secondary Color", "Hex Input / Colour Picker", "Complementary colour for accents and borders.")
note_box(doc, "Use your brand's exact hex code for a seamless match. If you don't know it, check your website's CSS.", "tip")

heading(doc, "Widget Width", 2)
para(doc, "Controls the maximum width of the combo builder on desktop screens.")
table(doc,
    ["Option", "Width", "Best For"],
    [
        ["Full Width", "100%",    "Full-bleed layouts, edge-to-edge themes"],
        ["Narrow",     "860 px",  "Minimalist themes with centred content"],
        ["Default",    "1140 px", "Most standard Shopify themes (recommended)"],
        ["Wide",       "1400 px", "Wide-format or magazine-style themes"],
        ["Full HD",    "1920 px", "Very wide screens or immersive layouts"],
    ],
    [1.3, 1.2, 3.7]
)
field_row(doc, "Custom Width", "Number (px)", "Enter any pixel value to override the presets.")
note_box(doc, "On mobile, the widget always stretches to full screen width regardless of this setting.", "tip")

heading(doc, "Product Grid", 2)
para(doc, "Controls how many product cards appear in each row of the selection grid on desktop.")
table(doc,
    ["Option", "Cards Per Row", "Best For"],
    [
        ["3 per row", "3", "Large product images, minimal text"],
        ["4 per row", "4", "Standard catalogue — most stores (default)"],
        ["5 per row", "5", "Stores with many products and compact cards"],
        ["6 per row", "6", "Very wide layouts with lots of products"],
    ],
    [1.3, 1.5, 3.4]
)

heading(doc, "Display Options", 2)
field_row(doc, "Show Savings Badge",       "Toggle", "Show a badge displaying how much the customer saves vs buying individually. Best with Fixed Price bundles.", default="Off")
field_row(doc, "Show Product Prices",      "Toggle", "Display individual product prices beneath each card. Recommended with Dynamic Price.", default="Off")
field_row(doc, "Show Out-of-Stock Products","Toggle", "Show OOS products greyed out. Off = hide OOS completely.", default="Off")

heading(doc, "Quick Reference", 2)
table(doc,
    ["Setting", "Recommended Default"],
    [
        ["Primary / Secondary Colour",  "Your brand hex code"],
        ["Widget Width",                "Default (1140 px)"],
        ["Product Grid",                "4 per row"],
        ["Show Savings Badge",          "On (if using Fixed Price below market value)"],
        ["Show Product Prices",         "On (if using Dynamic Price)"],
        ["Show Out-of-Stock",           "Off (for a cleaner grid)"],
    ],
    [2.5, 3.7]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  PAGE 8 — FAQ & TROUBLESHOOTING
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "Page 8 — FAQ & Troubleshooting", 1)
para(doc, "Quickly find answers to common questions and fix issues without raising a support ticket.", italic=True, color=GRAY)

heading(doc, "General Questions", 2)

heading(doc, "Q: What is the difference between a Simple Box and a Specific Box?", 3)
para(doc, "Simple Box = one product grid, customer picks any N items from your whole store or a category. "
         "Specific Box = multiple guided steps, each showing a different set of products. Best for structured bundles.")

heading(doc, "Q: How many boxes can I create?", 3)
para(doc, "The number of active boxes and monthly bundle orders depends on your plan. "
         "Check your current limit in the Dashboard → Order Credit Left card.")

heading(doc, "Q: Which plan should I choose?", 3)
table(doc,
    ["Plan", "Best For", "Monthly Orders"],
    [
        ["Free",    "Testing / just getting started",            "10"],
        ["Basic",   "Small stores just launching bundles",       "50"],
        ["Advance", "Growing stores with consistent bundle sales","100"],
        ["Plus",    "High-volume stores",                        "Unlimited"],
    ],
    [1.2, 2.8, 1.8]
)

heading(doc, "Installation & Setup", 2)

heading(doc, "Q: The widget is not showing on my storefront. What do I check?", 3)
para(doc, "Follow this checklist in order:")
for s in [
    "App Embed is On — Dashboard → Theme App Embed Status must show green 'On'.",
    "Block is placed — Combo Builder block must be added in the Shopify Theme Editor.",
    "Box is Active — Go to Manage Boxes and confirm the box Status is On.",
    "You are on the right page — widget appears where the block was placed, not everywhere.",
    "Hard-refresh the page — press Ctrl+Shift+R (Windows) or Cmd+Shift+R (Mac).",
]:
    para(doc, s, bullet=True, size=10.5)

heading(doc, "Q: I enabled the app embed but the status card still shows 'Off'. What do I do?", 3)
for s in [
    "Make sure you clicked Save in the Shopify Theme Editor after toggling the embed On.",
    "Refresh the MixBox Dashboard page.",
    "If still Off, try disabling and re-enabling the embed in the theme editor.",
]:
    para(doc, s, bullet=True, size=10.5)

heading(doc, "Box Configuration", 2)

heading(doc, "Q: Can I change the box after publishing it?", 3)
para(doc, "Yes. Go to Manage Boxes, find the box, and click Edit. "
         "All settings can be changed and saved at any time. Changes take effect immediately.")

heading(doc, "Q: What happens if I delete a box?", 3)
para(doc, "The box is permanently removed and the associated Shopify product is also deleted. "
         "This cannot be undone. If you want to hide a box temporarily, use Deactivate instead.")
note_box(doc, "Use Deactivate to hide a box temporarily. Reserve Delete for boxes you are certain you no longer need.", "warning")

heading(doc, "Q: What does 'Dynamic Price' mean?", 3)
para(doc, "Dynamic Price means the bundle total equals the sum of individual product prices selected. "
         "The price updates live on the storefront as the customer adds items. "
         "This is different from Fixed Price where the total is always the same.")

heading(doc, "Q: Can I add a discount to a Dynamic Price bundle?", 3)
para(doc, "Yes — but only on Specific Boxes. When you choose Dynamic Price on a Specific Box, "
         "a Discount Type dropdown appears. You can apply a percentage (e.g. 10% off) or "
         "a fixed amount discount (e.g. ₹100 off).")

heading(doc, "Q: How do I limit a step to show only certain products?", 3)
para(doc, "In the Steps Configuration section, set Step Scope to 'Select Collections' and click "
         "'Choose Step Collections' to pick one or more collections. Only products in those "
         "collections will appear for that step.")

heading(doc, "Analytics", 2)

heading(doc, "Q: My recent orders are not showing in Analytics. What do I do?", 3)
para(doc, "Click the '↻ Sync Orders' button at the top of the Analytics page. "
         "This forces a manual sync with Shopify.")

heading(doc, "Q: Why does Analytics show ₹0 even though I have orders?", 3)
for s in [
    "Check the date range — make sure the period includes the order dates.",
    "Check the Type Filter — switch to 'All Box' if orders may be from a different type.",
    "Try clicking Sync Orders.",
]:
    para(doc, s, bullet=True, size=10.5)

heading(doc, "Q: What is Conversion Rate?", 3)
para(doc, "Conversion Rate = (customers who completed a bundle order) ÷ "
         "(customers who visited a page with the widget) × 100. "
         "A typical good range for bundle widgets is 3–8%.")

heading(doc, "Billing & Orders", 2)

heading(doc, "Q: What happens when I reach my monthly order limit?", 3)
para(doc, "New bundle orders will not be processed until the limit resets at the start of your next billing month, "
         "or until you upgrade. A warning banner appears at 80% and again at 100% of your limit.")
note_box(doc, "Upgrade your plan before hitting 100% to avoid missing orders.", "important")

heading(doc, "Q: Do cancelled or refunded orders count toward the limit?", 3)
para(doc, "No — only completed orders count. Cancelled or refunded orders do not reduce your credit.")

# ── Footer ─────────────────────────────────────────────────────────────────────
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(100)
r = p.add_run("MixBox – Box & Bundle Builder")
_font(r, size=14, bold=True, color=GREEN)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Knowledge Base  |  Version 1.0  |  April 2026")
_font(r2, size=10, color=GRAY)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Prepared by Pryxotech  |  balvant@pryxotech.com")
_font(r3, size=10, italic=True, color=GRAY)

# ── Save ───────────────────────────────────────────────────────────────────────
doc.save(OUT_PATH)
print(f"[OK] Saved: {OUT_PATH}")
ss_files = os.listdir(SS_DIR) if os.path.exists(SS_DIR) else []
missing  = [f for f in [
    "01-dashboard.png","02-growth-apps.png","03-choose-bundle-type.png",
    "04-create-simple-box.png","05-create-specific-box.png",
    "06-steps-configuration.png","07-manage-boxes.png",
    "08-analytics.png","09-analytics-charts.png","10-widget-settings.png",
] if f not in ss_files]
if missing:
    print(f"\n[WARN] {len(missing)} screenshot(s) not found -- placeholders used:")
    for f in missing:
        print(f"   -> docs/knowledge-base/screenshots/{f}")
    print("\nPlace the screenshots in the folder above and re-run to embed them.")
else:
    print("[OK] All 10 screenshots embedded.")
